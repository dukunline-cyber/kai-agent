#!/usr/bin/env python3
"""
document_formatter.py — Engine generate dokumen akademik format DOCX.

Standar: Indonesian academic formatting (pedoman umum universitas Indonesia).
- Margin: 4-4-3-3 cm (kiri-atas-kanan-bawah)
- Font: Times New Roman 12pt (default), bisa diubah
- Spacing: 1.5 lines (default), bisa diubah
- Paragraf: first-line indent 1.27cm, justified
- Penomoran: BAB I, 1.1, 1.1.1, a., (1)
- Tabel: caption di atas (italic), note di bawah (italic), border horizontal
- Gambar: caption di bawah (italic), centered
- Persamaan: centered, nomor di kanan

Usage CLI:
  python3 document_formatter.py create --title "Judul" --author "Nama" --output skripsi.docx
  python3 document_formatter.py add-chapter --file skripsi.docx --number 1 --title "Pendahuluan"
  python3 document_formatter.py add-section --file skripsi.docx --number 1.1 --title "Latar Belakang"
  python3 document_formatter.py add-paragraph --file skripsi.docx --text "Isi paragraf..."
  python3 document_formatter.py add-table --file skripsi.docx --caption "Data Responden" --data '[["Nama","Usia"],["Andi",25]]'
  python3 document_formatter.py info --file skripsi.docx

Usage import:
  from document_formatter import AcademicDoc
  doc = AcademicDoc("Judul", "Nama", "Universitas X")
  doc.add_chapter(1, "Pendahuluan")
  doc.add_section("1.1", "Latar Belakang")
  doc.add_paragraph("Isi paragraf...")
  doc.add_table([["Nama","Usia"],["Andi",25]], caption="Data Responden", note="Sumber: primer")
  doc.save("skripsi.docx")
"""

import sys
import os
import json
import argparse
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("ERROR: python-docx not installed. Install with: pip install python-docx")
    sys.exit(1)


class AcademicDoc:
    """Dokumen akademik dengan formatting standar Indonesia."""

    # Default settings
    DEFAULT_FONT = "Times New Roman"
    DEFAULT_FONT_SIZE = 12
    DEFAULT_MARGIN = {"left": 4, "top": 4, "right": 3, "bottom": 3}  # cm
    DEFAULT_SPACING = 1.5
    DEFAULT_INDENT = 1.27  # cm (0.5 inch)

    def __init__(self, title: str = "", author: str = "", university: str = "",
                 program: str = "", font_name: str = None, font_size: int = None,
                 spacing: float = None, margins: dict = None):
        self.doc = Document()
        self.title = title
        self.author = author
        self.university = university
        self.program = program
        self.font_name = font_name or self.DEFAULT_FONT
        self.font_size = font_size or self.DEFAULT_FONT_SIZE
        self.spacing = spacing or self.DEFAULT_SPACING
        self.margins = margins or self.DEFAULT_MARGIN
        self._table_counter = 0
        self._figure_counter = 0
        self._equation_counter = 0
        self._current_chapter = 0

        self._setup_page()
        self._setup_styles()

    def _setup_page(self):
        """Set margin halaman."""
        for section in self.doc.sections:
            section.left_margin = Cm(self.margins["left"])
            section.top_margin = Cm(self.margins["top"])
            section.right_margin = Cm(self.margins["right"])
            section.bottom_margin = Cm(self.margins["bottom"])

    def _setup_styles(self):
        """Set default font dan spacing untuk seluruh dokumen."""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = self.font_name
        font.size = Pt(self.font_size)
        # Set East Asian font too (untuk kompatibilitas)
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), self.font_name)

        pf = style.paragraph_format
        pf.line_spacing = self.spacing
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

    def _set_cell_border(self, cell, **kwargs):
        """Set border untuk cell tabel.
        kwargs: top, bottom, left, right → dict with sz, val, color
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
        for edge, attrs in kwargs.items():
            element = parse_xml(
                f'<w:{edge} {nsdecls("w")} w:val="{attrs.get("val", "single")}" '
                f'w:sz="{attrs.get("sz", "4")}" w:space="0" '
                f'w:color="{attrs.get("color", "000000")}"/>'
            )
            tcBorders.append(element)
        tcPr.append(tcBorders)

    def _format_paragraph(self, para, indent=True, justify=True,
                          spacing=None, bold=False, italic=False,
                          alignment=None, font_size=None, color=None):
        """Apply formatting ke paragraph."""
        pf = para.paragraph_format
        if indent:
            pf.first_line_indent = Cm(self.DEFAULT_INDENT)
        if justify:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if alignment:
            para.alignment = alignment
        pf.line_spacing = spacing or self.spacing
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

        for run in para.runs:
            run.font.name = self.font_name
            run.font.size = Pt(font_size or self.font_size)
            run.bold = bold
            run.italic = italic
            if color:
                run.font.color.rgb = RGBColor(*color)
            # East Asian font
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:eastAsia'), self.font_name)

    # ========================
    # TITLE PAGE
    # ========================

    def add_title_page(self, title: str = None, author: str = None,
                       university: str = None, program: str = None, year: str = None):
        """Tambah halaman judul."""
        from datetime import datetime
        title = title or self.title
        author = author or self.author
        university = university or self.university
        program = program or self.program
        year = year or str(datetime.now().year)

        # Spacing atas
        for _ in range(3):
            self.doc.add_paragraph()

        # Judul (bold, centered, uppercase, 14pt)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title.upper())
        run.font.name = self.font_name
        run.font.size = Pt(14)
        run.bold = True
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), self.font_name)

        # Spacing
        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # Logo placeholder
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("[LOGO UNIVERSITAS]")
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)

        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # Author
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(author)
        run.font.name = self.font_name
        run.font.size = Pt(12)
        run.bold = True

        # NIM placeholder
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("[NIM]")
        run.font.size = Pt(12)

        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # Program studi
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(program)
        run.font.size = Pt(12)
        run.bold = True

        # Universitas
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(university.upper())
        run.font.size = Pt(14)
        run.bold = True

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(year)
        run.font.size = Pt(12)
        run.bold = True

        self.doc.add_page_break()

    # ========================
    # HEADINGS
    # ========================

    def add_chapter(self, number: int, title: str):
        """Tambah BAB heading (centered, bold, uppercase, 14pt, new page)."""
        self._current_chapter = number
        self._table_counter = 0
        self._figure_counter = 0
        self._equation_counter = 0

        self.doc.add_page_break()
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.space_before = Pt(0)
        pf.space_after = Pt(12)
        pf.line_spacing = self.spacing

        # "BAB I" in roman
        roman = self._to_roman(number)
        run = p.add_run(f"BAB {roman}")
        run.font.name = self.font_name
        run.font.size = Pt(14)
        run.bold = True
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), self.font_name)

        # Title
        p2 = self.doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf2 = p2.paragraph_format
        pf2.first_line_indent = Cm(0)
        pf2.space_after = Pt(12)
        pf2.line_spacing = self.spacing
        run2 = p2.add_run(title.upper())
        run2.font.name = self.font_name
        run2.font.size = Pt(14)
        run2.bold = True
        rPr2 = run2._element.get_or_add_rPr()
        rFonts2 = rPr2.get_or_add_rFonts()
        rFonts2.set(qn('w:eastAsia'), self.font_name)

    def add_section(self, number: str, title: str):
        """Tambah section heading (1.1, 1.2, dst) — bold, left-aligned, 12pt."""
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
        pf.line_spacing = self.spacing
        pf.keep_with_next = True

        run = p.add_run(f"{number}  {title}")
        run.font.name = self.font_name
        run.font.size = Pt(12)
        run.bold = True
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), self.font_name)

    def add_subsection(self, number: str, title: str):
        """Tambah subsection (1.1.1) — bold italic, left-aligned, 12pt."""
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.space_before = Pt(4)
        pf.space_after = Pt(4)
        pf.line_spacing = self.spacing
        pf.keep_with_next = True

        run = p.add_run(f"{number}  {title}")
        run.font.name = self.font_name
        run.font.size = Pt(12)
        run.bold = True
        run.italic = True
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), self.font_name)

    # ========================
    # BODY TEXT
    # ========================

    def add_paragraph(self, text: str, indent: bool = True, justify: bool = True,
                      spacing: float = None, bold: bool = False, italic: bool = False):
        """Tambah paragraf body text (first-line indent, justified, 1.5 spacing)."""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        self._format_paragraph(p, indent=indent, justify=justify,
                               spacing=spacing, bold=bold, italic=italic)
        return p

    def add_numbered_list(self, items: list, style: str = "number"):
        """Tambah numbered list.
        style: 'number' (1. 2. 3.), 'letter' (a. b. c.), 'paren' ((1) (2) (3))
        """
        for i, item in enumerate(items):
            p = self.doc.add_paragraph()
            pf = p.paragraph_format
            pf.first_line_indent = Cm(self.DEFAULT_INDENT)
            pf.left_indent = Cm(self.DEFAULT_INDENT)
            pf.line_spacing = self.spacing

            if style == "number":
                prefix = f"{i+1}. "
            elif style == "letter":
                prefix = f"{chr(97+i)}. "
            elif style == "paren":
                prefix = f"({i+1}) "
            else:
                prefix = f"{i+1}. "

            run = p.add_run(prefix + item)
            run.font.name = self.font_name
            run.font.size = Pt(self.font_size)
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:eastAsia'), self.font_name)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def add_quote(self, text: str, source: str = "", long_quote: bool = False):
        """Tambah quotation.
        Short quote (<40 words): inline dengan tanda kutip.
        Long quote (>=40 words): indent kiri 1.27cm, tanpa tanda kutip, italic.
        """
        if long_quote or len(text.split()) >= 40:
            p = self.doc.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent = Cm(self.DEFAULT_INDENT * 2)
            pf.right_indent = Cm(self.DEFAULT_INDENT)
            pf.first_line_indent = Cm(0)
            pf.line_spacing = 1.0  # single spacing for long quotes
            pf.space_before = Pt(6)
            pf.space_after = Pt(6)
            run = p.add_run(text)
            run.italic = True
            run.font.name = self.font_name
            run.font.size = Pt(self.font_size)
            if source:
                run2 = p.add_run(f" ({source})")
                run2.italic = True
                run2.font.size = Pt(self.font_size)
        else:
            p = self.doc.add_paragraph()
            run = p.add_run(f'"{text}"')
            if source:
                run2 = p.add_run(f" ({source})")
            self._format_paragraph(p)

    # ========================
    # TABLES
    # ========================

    def add_table(self, data: list, caption: str = "", note: str = "",
                  chapter_num: int = None, table_num: int = None,
                  style: str = "academic"):
        """Tambah tabel dengan caption dan note.
        data: list of lists (first row = header)
        caption: judul tabel (di atas, italic, centered)
        note: catatan/sumber (di bawah, italic, smaller)
        chapter_num: nomor bab untuk numbering (auto jika None)
        table_num: nomor urut tabel (auto jika None)
        style: 'academic' (horizontal borders only) or 'full' (all borders)
        """
        chap = chapter_num or self._current_chapter
        if table_num is not None:
            tnum = table_num
        else:
            self._table_counter += 1
            tnum = self._table_counter

        # Caption di atas tabel
        if caption:
            cap_p = self.doc.add_paragraph()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = cap_p.paragraph_format
            pf.first_line_indent = Cm(0)
            pf.space_before = Pt(6)
            pf.space_after = Pt(4)
            pf.line_spacing = self.spacing
            cap_text = f"Tabel {chap}.{tnum}  {caption}"
            run = cap_p.add_run(cap_text)
            run.italic = True
            run.font.name = self.font_name
            run.font.size = Pt(self.font_size)
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:eastAsia'), self.font_name)

        # Create table
        rows = len(data)
        cols = len(data[0]) if data else 0
        table = self.doc.add_table(rows=rows, cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Fill data
        for i, row_data in enumerate(data):
            for j, cell_data in enumerate(row_data):
                cell = table.cell(i, j)
                cell.text = str(cell_data)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
                pf = p.paragraph_format
                pf.first_line_indent = Cm(0)
                pf.line_spacing = 1.0

                for run in p.runs:
                    run.font.name = self.font_name
                    run.font.size = Pt(self.font_size)
                    rPr = run._element.get_or_add_rPr()
                    rFonts = rPr.get_or_add_rFonts()
                    rFonts.set(qn('w:eastAsia'), self.font_name)
                    if i == 0:  # header row
                        run.bold = True

        # Apply borders
        if style == "academic":
            # Horizontal borders only (APA style)
            for j in range(cols):
                # Top border (above header)
                self._set_cell_border(table.cell(0, j),
                    top={"sz": "8", "val": "single"},
                    bottom={"sz": "4", "val": "single"})
                # Bottom border (below header)
                if rows > 1:
                    self._set_cell_border(table.cell(0, j),
                        bottom={"sz": "4", "val": "single"})
                # Bottom border (below last row)
                self._set_cell_border(table.cell(rows-1, j),
                    bottom={"sz": "8", "val": "single"})
        else:
            # Full borders
            for i in range(rows):
                for j in range(cols):
                    self._set_cell_border(table.cell(i, j),
                        top={"sz": "4", "val": "single"},
                        bottom={"sz": "4", "val": "single"},
                        left={"sz": "4", "val": "single"},
                        right={"sz": "4", "val": "single"})

        # Note di bawah tabel
        if note:
            note_p = self.doc.add_paragraph()
            note_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf = note_p.paragraph_format
            pf.first_line_indent = Cm(0)
            pf.space_before = Pt(2)
            pf.space_after = Pt(6)
            pf.line_spacing = 1.0
            run = note_p.add_run(note)
            run.italic = True
            run.font.name = self.font_name
            run.font.size = Pt(10)  # smaller font for note
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:eastAsia'), self.font_name)
        else:
            # Add spacing after table
            self.doc.add_paragraph()

    # ========================
    # FIGURES
    # ========================

    def add_figure(self, image_path: str, caption: str = "",
                   chapter_num: int = None, figure_num: int = None,
                   width_cm: float = 12):
        """Tambah gambar dengan caption di bawah."""
        chap = chapter_num or self._current_chapter
        if figure_num is not None:
            fnum = figure_num
        else:
            self._figure_counter += 1
            fnum = self._figure_counter

        # Image (centered)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.space_before = Pt(6)
        pf.space_after = Pt(4)
        run = p.add_run()
        if os.path.exists(image_path):
            run.add_picture(image_path, width=Cm(width_cm))
        else:
            run.add_text(f"[IMAGE: {image_path}]")
            run.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)

        # Caption di bawah (italic, centered)
        if caption:
            cap_p = self.doc.add_paragraph()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = cap_p.paragraph_format
            pf.first_line_indent = Cm(0)
            pf.space_before = Pt(2)
            pf.space_after = Pt(6)
            pf.line_spacing = self.spacing
            cap_text = f"Gambar {chap}.{fnum}  {caption}"
            run = cap_p.add_run(cap_text)
            run.italic = True
            run.font.name = self.font_name
            run.font.size = Pt(self.font_size)
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:eastAsia'), self.font_name)

    # ========================
    # EQUATIONS
    # ========================

    def add_equation(self, equation: str, chapter_num: int = None,
                     equation_num: int = None):
        """Tambah persamaan (centered, nomor di kanan)."""
        chap = chapter_num or self._current_chapter
        if equation_num is not None:
            enum = equation_num
        else:
            self._equation_counter += 1
            enum = self._equation_counter

        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
        pf.line_spacing = self.spacing

        # Tab stops: center for equation, right for number
        from docx.enum.text import WD_TAB_ALIGNMENT
        pf.tab_stops.add_tab_stop(Cm(7.5), WD_TAB_ALIGNMENT.CENTER)
        pf.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.RIGHT)

        run = p.add_run(f"\t{equation}\t({chap}.{enum})")
        run.font.name = self.font_name
        run.font.size = Pt(self.font_size)
        run.italic = True
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), self.font_name)

    # ========================
    # REFERENCES
    # ========================

    def add_references_heading(self):
        """Tambah heading Daftar Pustaka (new page, centered, bold)."""
        self.doc.add_page_break()
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.space_after = Pt(12)
        pf.line_spacing = self.spacing
        run = p.add_run("DAFTAR PUSTAKA")
        run.font.name = self.font_name
        run.font.size = Pt(14)
        run.bold = True
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), self.font_name)

    def add_reference_entry(self, entry: str, hanging_indent: bool = True):
        """Tambah entry daftar pustaka dengan hanging indent."""
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.first_line_indent = Cm(-self.DEFAULT_INDENT)  # hanging indent
        pf.left_indent = Cm(self.DEFAULT_INDENT)
        pf.line_spacing = self.spacing
        pf.space_after = Pt(6)
        run = p.add_run(entry)
        run.font.name = self.font_name
        run.font.size = Pt(self.font_size)
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), self.font_name)

    # ========================
    # UTILITIES
    # ========================

    def add_page_break(self):
        """Tambah page break."""
        self.doc.add_page_break()

    def add_blank_line(self, count: int = 1):
        """Tambah baris kosong."""
        for _ in range(count):
            self.doc.add_paragraph()

    def add_toc_placeholder(self):
        """Tambah placeholder Table of Contents."""
        self.doc.add_page_break()
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.space_after = Pt(12)
        run = p.add_run("DAFTAR ISI")
        run.font.name = self.font_name
        run.font.size = Pt(14)
        run.bold = True
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), self.font_name)

        p2 = self.doc.add_paragraph()
        run2 = p2.add_run("[Generate Daftar Isi via Word: References → Table of Contents]")
        run2.italic = True
        run2.font.color.rgb = RGBColor(128, 128, 128)
        run2.font.size = Pt(10)

    @staticmethod
    def _to_roman(num: int) -> str:
        """Convert integer to Roman numeral."""
        val = [
            (1000, "M"), (900, "CM"), (500, "D"), (400, "CD"),
            (100, "C"), (90, "XC"), (50, "L"), (40, "XL"),
            (10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I")
        ]
        result = ""
        for v, s in val:
            while num >= v:
                result += s
                num -= v
        return result

    def save(self, filepath: str):
        """Simpan dokumen."""
        self.doc.save(filepath)
        return filepath

    def info(self) -> dict:
        """Return info tentang dokumen."""
        para_count = len(self.doc.paragraphs)
        table_count = len(self.doc.tables)
        word_count = sum(len(p.text.split()) for p in self.doc.paragraphs)
        return {
            "paragraphs": para_count,
            "tables": table_count,
            "words": word_count,
            "font": self.font_name,
            "font_size": self.font_size,
            "spacing": self.spacing,
            "margins": self.margins,
        }


# ========================
# CLI
# ========================

def cli():
    parser = argparse.ArgumentParser(description="Academic Document Formatter")
    sub = parser.add_subparsers(dest="action")

    # Create
    p_create = sub.add_parser("create", help="Create new document")
    p_create.add_argument("--title", required=True)
    p_create.add_argument("--author", default="")
    p_create.add_argument("--university", default="")
    p_create.add_argument("--program", default="")
    p_create.add_argument("--output", required=True)
    p_create.add_argument("--font", default="Times New Roman")
    p_create.add_argument("--font-size", type=int, default=12)
    p_create.add_argument("--spacing", type=float, default=1.5)
    p_create.add_argument("--title-page", action="store_true")

    # Add chapter
    p_chap = sub.add_parser("add-chapter", help="Add chapter")
    p_chap.add_argument("--file", required=True)
    p_chap.add_argument("--number", type=int, required=True)
    p_chap.add_argument("--title", required=True)

    # Add section
    p_sec = sub.add_parser("add-section", help="Add section")
    p_sec.add_argument("--file", required=True)
    p_sec.add_argument("--number", required=True)
    p_sec.add_argument("--title", required=True)

    # Add subsection
    p_sub = sub.add_parser("add-subsection", help="Add subsection")
    p_sub.add_argument("--file", required=True)
    p_sub.add_argument("--number", required=True)
    p_sub.add_argument("--title", required=True)

    # Add paragraph
    p_para = sub.add_parser("add-paragraph", help="Add paragraph")
    p_para.add_argument("--file", required=True)
    p_para.add_argument("--text", required=True)
    p_para.add_argument("--no-indent", action="store_true")

    # Add table
    p_table = sub.add_parser("add-table", help="Add table")
    p_table.add_argument("--file", required=True)
    p_table.add_argument("--caption", default="")
    p_table.add_argument("--note", default="")
    p_table.add_argument("--data", required=True, help='JSON: [["h1","h2"],["r1c1","r1c2"]]')

    # Add figure
    p_fig = sub.add_parser("add-figure", help="Add figure")
    p_fig.add_argument("--file", required=True)
    p_fig.add_argument("--image", required=True)
    p_fig.add_argument("--caption", default="")

    # Add reference
    p_ref = sub.add_parser("add-reference", help="Add reference entry")
    p_ref.add_argument("--file", required=True)
    p_ref.add_argument("--entry", required=True)
    p_ref.add_argument("--heading", action="store_true")

    # Info
    p_info = sub.add_parser("info", help="Get document info")
    p_info.add_argument("--file", required=True)

    args = parser.parse_args()

    if args.action == "create":
        doc = AcademicDoc(
            title=args.title, author=args.author,
            university=args.university, program=args.program,
            font_name=args.font, font_size=args.font_size,
            spacing=args.spacing
        )
        if args.title_page:
            doc.add_title_page()
        doc.save(args.output)
        print(f"Created: {args.output}")

    elif args.action == "add-chapter":
        doc = Document(args.file)
        # Re-wrap using AcademicDoc but keep existing doc
        ad = AcademicDoc()
        ad.doc = doc
        ad.add_chapter(args.number, args.title)
        ad.save(args.file)
        print(f"Added chapter {args.number}: {args.title}")

    elif args.action == "add-section":
        doc = Document(args.file)
        ad = AcademicDoc()
        ad.doc = doc
        ad.add_section(args.number, args.title)
        ad.save(args.file)
        print(f"Added section {args.number}: {args.title}")

    elif args.action == "add-subsection":
        doc = Document(args.file)
        ad = AcademicDoc()
        ad.doc = doc
        ad.add_subsection(args.number, args.title)
        ad.save(args.file)
        print(f"Added subsection {args.number}: {args.title}")

    elif args.action == "add-paragraph":
        doc = Document(args.file)
        ad = AcademicDoc()
        ad.doc = doc
        ad.add_paragraph(args.text, indent=not args.no_indent)
        ad.save(args.file)
        print(f"Added paragraph ({len(args.text)} chars)")

    elif args.action == "add-table":
        doc = Document(args.file)
        ad = AcademicDoc()
        ad.doc = doc
        data = json.loads(args.data)
        ad.add_table(data, caption=args.caption, note=args.note)
        ad.save(args.file)
        print(f"Added table: {args.caption}")

    elif args.action == "add-figure":
        doc = Document(args.file)
        ad = AcademicDoc()
        ad.doc = doc
        ad.add_figure(args.image, caption=args.caption)
        ad.save(args.file)
        print(f"Added figure: {args.caption}")

    elif args.action == "add-reference":
        doc = Document(args.file)
        ad = AcademicDoc()
        ad.doc = doc
        if args.heading:
            ad.add_references_heading()
        ad.add_reference_entry(args.entry)
        ad.save(args.file)
        print(f"Added reference entry")

    elif args.action == "info":
        doc = Document(args.file)
        ad = AcademicDoc()
        ad.doc = doc
        info = ad.info()
        print(json.dumps(info, indent=2))

    else:
        parser.print_help()


if __name__ == "__main__":
    cli()
